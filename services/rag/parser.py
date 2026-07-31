"""
Document parser — converts raw file bytes into structured page text.

Supports: PDF (via PyMuPDF), DOCX (python-docx), HTML (BeautifulSoup),
plain text. Extracts text with page/section metadata for chunk attribution.

Graceful degradation: unknown formats fall back to raw UTF-8 decode.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TEXT = "text"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"


@dataclass
class ParsedPage:
    """A single page or section extracted from a document."""
    page_num: int
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """A fully parsed document ready for chunking."""
    title: str
    url: str
    format: DocumentFormat
    pages: list[ParsedPage]
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def total_chars(self) -> int:
        return sum(len(p.text) for p in self.pages)


class DocumentParser:
    """
    Multi-format document parser.

    Usage:
        parser = DocumentParser()
        parsed = parser.parse(raw_bytes, url="https://example.com/doc.pdf")
    """

    def parse(
        self,
        data: bytes | BinaryIO,
        *,
        url: str = "",
        title: str = "",
        hint_format: DocumentFormat | None = None,
    ) -> ParsedDocument:
        """
        Parse document bytes into a ParsedDocument.

        Args:
            data:        Raw bytes or file-like object.
            url:         Source URL for metadata.
            title:       Document title hint.
            hint_format: If known, skip format detection.

        Returns:
            ParsedDocument with pages populated.
        """
        if isinstance(data, (bytes, bytearray)):
            buf = io.BytesIO(data)
        else:
            buf = data

        fmt = hint_format or self._detect_format(url, buf)
        buf.seek(0)

        try:
            if fmt == DocumentFormat.PDF:
                pages = self._parse_pdf(buf)
            elif fmt == DocumentFormat.DOCX:
                pages = self._parse_docx(buf)
            elif fmt in (DocumentFormat.HTML,):
                pages = self._parse_html(buf)
            elif fmt in (DocumentFormat.TEXT, DocumentFormat.MARKDOWN):
                pages = self._parse_text(buf)
            else:
                pages = self._parse_text(buf)
        except Exception as exc:
            logger.warning("Parser failed for %r (%s): %s — falling back to text.", url, fmt, exc)
            buf.seek(0)
            pages = self._parse_text(buf)

        return ParsedDocument(
            title=title or self._extract_title(pages, url),
            url=url,
            format=fmt,
            pages=pages,
            metadata={"url": url, "format": fmt.value},
        )

    # ─── Format detection ─────────────────────────────────────────────

    def _detect_format(self, url: str, buf: io.BytesIO) -> DocumentFormat:
        """Detect document format from URL extension then magic bytes."""
        suffix = Path(url.split("?")[0]).suffix.lower()
        if suffix == ".pdf":
            return DocumentFormat.PDF
        if suffix in (".docx", ".doc"):
            return DocumentFormat.DOCX
        if suffix in (".htm", ".html"):
            return DocumentFormat.HTML
        if suffix in (".md", ".markdown"):
            return DocumentFormat.MARKDOWN
        if suffix in (".txt",):
            return DocumentFormat.TEXT

        # Magic bytes
        magic = buf.read(8)
        buf.seek(0)
        if magic[:4] == b"%PDF":
            return DocumentFormat.PDF
        if magic[:2] == b"PK":  # DOCX is a ZIP
            return DocumentFormat.DOCX
        if b"<html" in magic.lower() or b"<!doctype" in magic.lower():
            return DocumentFormat.HTML

        return DocumentFormat.TEXT

    # ─── Format-specific parsers ──────────────────────────────────────

    def _parse_pdf(self, buf: io.BytesIO) -> list[ParsedPage]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            return [ParsedPage(page_num=0, text=buf.read().decode("utf-8", errors="replace"))]

        pages = []
        with fitz.open(stream=buf, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    pages.append(ParsedPage(
                        page_num=page_num,
                        text=text,
                        metadata={"width": page.rect.width, "height": page.rect.height},
                    ))
        return pages

    def _parse_docx(self, buf: io.BytesIO) -> list[ParsedPage]:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return [ParsedPage(page_num=0, text=buf.read().decode("utf-8", errors="replace"))]

        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [ParsedPage(page_num=0, text=full_text)]

    def _parse_html(self, buf: io.BytesIO) -> list[ParsedPage]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("BeautifulSoup not installed. Run: pip install beautifulsoup4")
            return [ParsedPage(page_num=0, text=buf.read().decode("utf-8", errors="replace"))]

        html = buf.read().decode("utf-8", errors="replace")
        soup = BeautifulSoup(html, "lxml")

        # Remove boilerplate tags
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        title = soup.title.string if soup.title else ""
        return [ParsedPage(page_num=0, text=text, metadata={"html_title": title})]

    def _parse_text(self, buf: io.BytesIO) -> list[ParsedPage]:
        text = buf.read().decode("utf-8", errors="replace")
        return [ParsedPage(page_num=0, text=text)]

    # ─── Helpers ──────────────────────────────────────────────────────

    def _extract_title(self, pages: list[ParsedPage], url: str) -> str:
        """Best-effort title extraction from page metadata or first line."""
        for page in pages:
            if page.metadata.get("html_title"):
                return page.metadata["html_title"]
        if pages and pages[0].text:
            first_line = pages[0].text.strip().split("\n")[0]
            if len(first_line) < 200:
                return first_line
        return url.split("/")[-1] or "Untitled"
